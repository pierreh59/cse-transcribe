# -*- coding: utf-8 -*-
"""
Exporte une transcription diarisee (transcript_diarized.json) en remplacant les
identifiants de locuteur (SPEAKER_00, ...) par un nom/prenom/fonction fournis par
l'utilisateur, au format Word (.docx), PDF (.pdf) ou texte (.txt).

Invoque comme sous-processus par l'interface graphique (SpeakerNamingDialog) apres
validation de l'identification des locuteurs.
"""
import argparse
import json
import os


def fmt_hms(t: float) -> str:
    h = int(t // 3600)
    m = int((t % 3600) // 60)
    s = int(t % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def build_label(speaker_id: str, mapping: dict) -> str:
    info = mapping.get(speaker_id)
    if info:
        prenom = (info.get("prenom") or "").strip()
        nom = (info.get("nom") or "").strip()
        fonction = (info.get("fonction") or "").strip()
        full = f"{prenom} {nom}".strip() or speaker_id
        return f"{full} ({fonction})" if fonction else full
    return "Intervenant non identifié"


def write_txt(turns, path):
    with open(path, "w", encoding="utf-8") as f:
        for t in turns:
            f.write(f"[{fmt_hms(t['start'])}] {t['speaker_label']} : {t['text'].strip()}\n")


def write_docx(turns, path, title):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()
    for t in turns:
        p = doc.add_paragraph()
        r = p.add_run(f"[{fmt_hms(t['start'])}] {t['speaker_label']} : ")
        r.bold = True
        p.add_run(t["text"].strip())
    doc.save(path)


def write_pdf(turns, path, title):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Les polices de base de fpdf2 (Helvetica, Times...) ne couvrent que le
    # Latin-1 : Whisper restitue souvent une ponctuation typographique
    # (apostrophes courbes, points de suspension, tirets longs) hors de ce
    # jeu, ce qui fait planter le decoupage de lignes. On utilise Arial
    # (present sur tout poste Windows), qui couvre correctement le francais.
    font_name = "Helvetica"
    windows_fonts = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
    regular = os.path.join(windows_fonts, "arial.ttf")
    bold = os.path.join(windows_fonts, "arialbd.ttf")
    if os.path.isfile(regular) and os.path.isfile(bold):
        pdf.add_font("Arial", "", regular)
        pdf.add_font("Arial", "B", bold)
        font_name = "Arial"

    # multi_cell(w=0, ...) laisse par defaut le curseur X au bord droit de la
    # page une fois la cellule ecrite (au lieu de revenir a la marge gauche) :
    # l'appel suivant se retrouve alors avec une largeur disponible nulle et
    # fpdf2 leve une exception des le premier caractere. On force explicitement
    # le retour a la marge gauche apres chaque cellule.
    cell_kwargs = {"new_x": "LMARGIN", "new_y": "NEXT"}

    pdf.set_font(font_name, "B", 14)
    pdf.multi_cell(0, 10, title, **cell_kwargs)
    pdf.ln(4)
    for t in turns:
        pdf.set_font(font_name, "B", 10)
        pdf.multi_cell(0, 6, f"[{fmt_hms(t['start'])}] {t['speaker_label']} :", **cell_kwargs)
        pdf.set_font(font_name, size=10)
        text = t["text"].strip() or " "
        pdf.multi_cell(0, 6, text, **cell_kwargs)
        pdf.ln(1)
    pdf.output(path)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--transcript", required=True, help="Chemin vers transcript_diarized.json")
    parser.add_argument("--mapping", required=True,
                         help="Fichier JSON {speaker_id: {nom, prenom, fonction}}")
    parser.add_argument("--format", required=True, choices=["docx", "pdf", "txt"])
    parser.add_argument("--output", required=True)
    parser.add_argument("--title", default="Transcription")
    args = parser.parse_args()

    with open(args.transcript, encoding="utf-8") as f:
        turns = json.load(f)
    with open(args.mapping, encoding="utf-8") as f:
        mapping = json.load(f)

    for t in turns:
        t["speaker_label"] = build_label(t["speaker"], mapping)

    if args.format == "docx":
        write_docx(turns, args.output, args.title)
    elif args.format == "pdf":
        write_pdf(turns, args.output, args.title)
    else:
        write_txt(turns, args.output)

    print(f"Export ecrit : {args.output}")


if __name__ == "__main__":
    main()
